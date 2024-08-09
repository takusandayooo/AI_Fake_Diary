import docx
import time
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches
from docx2pdf import convert

def text_split(text, max_length):
    # 文字列を指定した文字数で分割する
    return [text[i:i + max_length] for i in range(0, len(text), max_length)]
def make_word(result):
    doc = docx.Document("./word/sample.docx")

    tbl = doc.tables[0]

    # 日付情報を代入
    dateInfo = tbl.rows[1].cells[-1].text
    cell = tbl.rows[1].cells[-1]
    han_to_zen = str.maketrans("0123456789", "０１２３４５６７８９")
    dayOfWeek = {"Sunday":"日曜日","Monday":"月曜日","Tuesday":"火曜日","Wednesday":"水曜日","Thursday":"木曜日","Friday":"金曜日","Saturday":"土曜日"}
    dateInfo = dateInfo.replace("MM", str(int(time.strftime("%m"))))
    dateInfo = dateInfo.replace("DD", str(int(time.strftime("%d"))))
    dateInfo = dateInfo.replace("W", dayOfWeek[time.strftime("%A")])
    dateInfo = dateInfo.translate(han_to_zen)  # NOTE: 全角数字の方が見栄えが良いため変換している
    cell.text = dateInfo

    # フォーマットを変更
    paragraph = cell.paragraphs[0]
    run = paragraph.runs[0]
    run.font.size = Pt(16)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 絵日記に任意の文字列を入力
    cell = tbl.rows[1].cells[0:-1]
    splitText = text_split(result, 14)
    reversed_cells = list(reversed(cell))

    for selectCell, text in zip(reversed_cells, splitText):
        selectCell.text = text
        paragraph = selectCell.paragraphs[0]
        run = paragraph.runs[0]
        run.font.size = Pt(24)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # 写真を挿入
    cell = tbl.rows[0].cells[0]
    cell_width = cell.width
    table_width = cell_width
    table_height = cell_width*0.75

    # 画像の幅と高さをセルのサイズに合わせて調整
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run()
    run.add_picture('./photo/result.png', width=table_width, height=table_height)

    doc.save("./word/sample2.docx")
    convert("./word/sample2.docx", "./static/sample2.pdf")



